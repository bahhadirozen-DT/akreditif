import json
import os
import re
import sys
from datetime import datetime, timedelta
from docx import Document
from ucp600_motoru import UCP600KuralMotoru, normalize_text
from dokuman_yonetici import DokumanYonetici

def safe_float(val_str):
    """Finansal tutarları güvenli bir şekilde float tipine dönüştürür."""
    if not val_str:
        return 0.0
    val_str = val_str.strip()
    if (',' in val_str and '.' in val_str and val_str.rfind(',') > val_str.rfind('.')) or (',' in val_str and '.' not in val_str):
        val_str = val_str.replace('.', '').replace(',', '.')
    else:
        val_str = val_str.replace(',', '')
    
    val_str = re.sub(r'[^0-9.]', '', val_str)
    return float(val_str) if val_str else 0.0

def load_rules(config_path="kurallar.json"):
    """JSON tabanlı kuralları yükler."""
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Kritik Hata: {config_path} bulunamadı!")
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)

def evrak_sozluge_cevir(metin):
    """Satır bazlı ham metni sözlük yapısına dönüştürür."""
    sonuc = {}
    for satir in metin.split('\n'):
        if ':' in satir:
            a, d = satir.split(':', 1)
            anahtar = a.strip().upper().replace("İ", "I").replace(" ", "_")
            sonuc[anahtar] = d.strip()
    return sonuc

def madde_uzman_yorumu_al(madde_adi):
    """Bulunan UCP 600 maddelerine profesyonel dış ticaret yorumu ekler."""
    yorumlar = {
        "Art 2": "Tanımlamalar maddesidir. Akreditif süreçlerindeki tarafların (Amir, Lehtar, İhbar Bankası) sorumluluk sınırlarını çizer.",
        "Art 4": "Akreditiflerin dayandığı temel satış sözleşmelerinden tamamen bağımsız bir hukuki işlem olduğunu belirtir. Bankalar sadece belgelerle ilgilenir, malların fiziki durumuyla ilgilenmez.",
        "Art 5": "Bankaların işlemlerinin sadece belgeler üzerinden yürüdüğünü, mallar, hizmetler veya performanslarla ilgilenmediğini vurgular (Belge Ticareti İlkesi).",
        "Art 6": "Akreditifin hangi bankanın gişesinde (kullanım yerinde) son bulacağını ve ödeme şeklini (Görüldüğünde, Vadeli, İştira) netleştirir.",
        "Art 7": "Amir bankanın (Issuing Bank), kurallara uygun evrak ibraz edildiğinde lehtara karşı geri dönülemez ve kesin bir ödeme taahhüdü altında olduğunu açıklar.",
        "Art 8": "Teyit bankasının sorumluluklarını belirler. Teyit eklenmişse, amir bankanın riskine ek olarak ikinci bir bankanın ödeme garantisi devreye girer.",
        "Art 14": "Bankaların belgeleri inceleme standartlarını belirler. Belgelerin kendi arasında ve akreditifle çelişmemesi (tutarlılık) için en kritik maddedir.",
        "Art 18": "Ticari faturanın (Commercial Invoice) mutlaka lehtar tarafından düzenlenmesi, amir adına kesilmesi ve akreditif döviziyle uyumlu olması şartını koşar.",
        "Art 20": "Deniz Konşimentosu (Bill of Lading) kurallarıdır. Belgenin mutlaka 'Shipped on Board' (Yüklendi) kaydı taşıması, orijinal olması ve taşıyanın imzasını barındırması zorunludur.",
        "Art 27": "Taşıma belgesinin 'Temiz' (Clean) olması gerektiğini, yani mal veya ambalajda hasar/kusur belirten hiçbir ibare taşımaması gerektiğini söyler.",
        "Art 28": "Sigorta belgelerinin akreditif dövizinde olması, en geç yükleme gününde başlaması ve aksi belirtilmedikçe fatura tutarının en az %110'unu kapsaması şarttır.",
        "Art 31": "Kısmi yükleme (Partial Shipment) izinlerini düzenler. Akreditifte yasaklanmadığı sürece parça parça mal sevkiyatı yapmak serbesttir."
    }
    madde_temiz = madde_adi.replace("[", "").replace("]", "").strip()
    return yorumlar.get(madde_temiz, "UCP 600 kuralları gereği bu maddenin operasyonel olarak yönetilmesi rezerv risklerini azaltır.")

def markdown_raporu_olustur(t_not, z_sonuc, k_tespit, k_eksik, e_sonuc, f_not, i_not, ucp_sonuc, h_var):
    with open("akreditif_analiz_raporu.md", "w", encoding="utf-8") as f:
        f.write("# 📋 AKREDİTİF GELİŞMİŞ UZMAN DENETİM RAPORU\n\n")
        f.write(f"**Rapor Tarihi:** {datetime.now().strftime('%d.%m.%Y %H:%M')}\n---\n\n")
        
        if t_not:
            f.write("## 1. Kritik Süreler ve Vade Analizi\n")
            for n in t_not: f.write(f"* {n}\n")
            f.write("\n")
        if f_not:
            f.write("## 2. Finansal Vade ve Ödeme Takvimi\n")
            for n in f_not: f.write(f"* {n}\n")
            f.write("\n")
            
        f.write("## 3. Incoterms ve Sigorta Denetimi\n")
        for d, m in i_not:
            e = "✅" if "UYUMLU" in d or "BİLGİ" in d or "BILGI" in d else "🚨"
            f.write(f"* {e} **[{d}]** {m}\n")
        f.write("\n")
            
        if e_sonuc:
            f.write("## 4. Çapraz Evrak Uyumluluk Kontrolü\n")
            for d, m in e_sonuc:
                e = "✅" if "UYUMLU" in d else "🚨"
                f.write(f"* {e} **[{d}]** {m}\n")
            f.write("\n")
            
        if z_sonuc:
            f.write("## 5. Zorunlu UCP 600 Parametreleri\n")
            for d, m in z_sonuc:
                e = "✅" if d == "UYUMLU" else "❌"
                f.write(f"* {e} **[{d}]** {m}\n")
            f.write("\n")
            
        f.write("## 6. UCP 600 Maddeleri ve Uzman Yorum Tablosu\n")
        if k_tespit:
            for d, m in k_tespit:
                match = re.search(r'\[(.*?)\]', m)
                m_adi = match.group(1) if match else "UCP"
                yorum = madde_uzman_yorumu_al(m_adi)
                f.write(f"* 🔍 **[{d}]** {m} \n  * 💡 *Uzman Analizi:* {yorum}\n")
        if k_eksik:
            maddeler_str = ", ".join([m for m, _ in k_eksik])
            f.write(f"* ⚠ **[MANUEL KONTROL]** Metinde doğrudan geçmeyen maddeler: *{maddeler_str}*\n")
        f.write("\n")
            
        if ucp_sonuc:
            f.write("## 7. UCP 600 Resmi Kural Motoru Çıktıları\n")
            for d, m in ucp_sonuc:
                e = "✅" if "UYUMLU" in d or "BİLGİ" in d else "🚨"
                f.write(f"* {e} **[{d}]** {m}\n")
            f.write("\n")
            
        f.write("---\n")
        if h_var:
            f.write("### 🚨 SONUÇ: Rezerv riskleri tespit edildi! Kontrol etmeden bankaya vermeyin.\n")
        else:
            f.write("### 🎉 SONUÇ: Tüm kontroller başarıyla tamamlandı. Altyapı tam uyumlu.\n")

def word_raporu_olustur(t_not, z_sonuc, k_tespit, k_eksik, e_sonuc, f_not, i_not, ucp_sonuc, h_var):
    doc = Document()
    doc.add_heading('GELİŞMİŞ DIŞ TİCARET DENETİM VE ANALİZ RAPORU', level=1)
    doc.add_paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
    
    if t_not:
        doc.add_heading('1. Kritik Süreler ve Vade Analizi', level=2)
        for n in t_not: doc.add_paragraph(str(n))
    if f_not:
        doc.add_heading('2. Finansal Vade ve Ödeme Takvimi', level=2)
        for n in f_not: doc.add_paragraph(str(n))
        
    doc.add_heading('3. Incoterms ve Sigorta Denetimi', level=2)
    for d, m in i_not: doc.add_paragraph(f"[{str(d)}] {str(m)}")
        
    if e_sonuc:
        doc.add_heading('4. Çapraz Evrak Uyumluluk Kontrolü', level=2)
        for d, m in e_sonuc: doc.add_paragraph(f"[{str(d)}] {str(m)}")
        
    if z_sonuc:
        doc.add_heading('5. Zorunlu UCP 600 Parametreleri', level=2)
        for d, m in z_sonuc: doc.add_paragraph(f"[{str(d)}] {str(m)}")
        
    doc.add_heading('6. UCP 600 Denetim ve Mevzuat Yorum Tablosu', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kural / Madde'
    hdr_cells[1].text = 'Durum'
    hdr_cells[2].text = 'Bulgu Mesajı'
    hdr_cells[3].text = 'Hukuki ve Pratik Uzman Yorumu'
    
    for d, m in k_tespit:
        row_cells = table.add_row().cells
        match = re.search(r'\[(.*?)\]', m)
        m_adi = match.group(1) if match else "UCP"
        row_cells[0].text = m_adi
        row_cells[1].text = str(d)
        row_cells[2].text = m
        row_cells[3].text = madde_uzman_yorumu_al(m_adi)
        
    for d, m in ucp_sonuc:
        row_cells = table.add_row().cells
        match = re.search(r'(Madde \d+\(?[a-z]?\)?):', m)
        m_adi = match.group(1) if match else "UCP Motoru"
        row_cells[0].text = m_adi
        row_cells[1].text = str(d)
        row_cells[2].text = m
        row_cells[3].text = "Bu bulgu, akreditif kuralları (UCP 600) ve uluslararası bankacılık standart uygulamalarının (ISBP) çapraz evrak kontrolünden üretilmiştir."

    if k_eksik:
        doc.add_heading('7. Doğrudan Geçmeyen Maddeler Notu', level=2)
        maddeler_str = ", ".join([m for m, _ in k_eksik])
        doc.add_paragraph(f"Aşağıdaki kurallara ait spesifik ifadeler SWIFT metninde doğrudan yer almamaktadır:\n{maddeler_str}")
        
    doc.add_paragraph("\n" + "="*40)
    if h_var:
        doc.add_paragraph("SONUÇ: Rezerv riskleri tespit edildi! Evrakları bankaya ibraz etmeden önce riskli alanları revize edin.")
    else:
        doc.add_paragraph("SONUÇ: Altyapı tam uyumlu.")
        
    doc.save("akreditif_analiz_raporu.docx")

def dinamik_dosya_bul_ve_oku(yonetici, varsayilan_ad, desteklenen_uzantilar):
    """
    Klasörde ve 'yuklenen_dosyalar/' dizininde evrakları otomatik tarar.
    """
    klasor = "yuklenen_dosyalar"
    
    for uzanti in desteklenen_uzantilar:
        # Önce yeni klasörün içinde ara
        test_yolu_yeni = os.path.join(klasor, f"{varsayilan_ad}{uzanti}")
        if os.path.exists(test_yolu_yeni):
            print(f"[OKUNUYOR] Klasörden evrak algılandı: {test_yolu_yeni}")
            return yonetici.evrak_metne_cevir(test_yolu_yeni)
            
        # Geriye dönük uyumluluk için eski ana dizindeki yerini de kontrol et
        test_yolu_eski = f"{varsayilan_ad}{uzanti}"
        if os.path.exists(test_yolu_eski):
            print(f"[OKUNUYOR] Ana dizinden evrak algılandı: {test_yolu_eski}")
            return yonetici.evrak_metne_cevir(test_yolu_eski)
            
    # Düz metin (.txt) kontrolleri
    txt_yolu_yeni = os.path.join(klasor, f"{varsayilan_ad}.txt")
    if os.path.exists(txt_yolu_yeni):
        return yonetici.evrak_metne_cevir(txt_yolu_yeni)
        
    txt_yolu_eski = f"{varsayilan_ad}.txt"
    if os.path.exists(txt_yolu_eski):
        return yonetici.evrak_metne_cevir(txt_yolu_eski)
        
    return ""

def analiz_yurut():
    kural_dosyasi = 'kurRules.json' if os.path.exists('kurRules.json') else 'kurallar.json'
    try:
        kurallar = load_rules(kural_dosyasi)
    except FileNotFoundError as e:
        print(e)
        sys.exit(1)

    yonetici = DokumanYonetici()
    uzantilar = ['.docx', '.xlsx', '.pdf', '.png', '.jpg', '.jpeg']

    kusat_upper = dinamik_dosya_bul_ve_oku(yonetici, "gelen_kusat", uzantilar).upper()
    fatura_metin = dinamik_dosya_bul_ve_oku(yonetici, "fatura", uzantilar)
    konsimento_metin = dinamik_dosya_bul_ve_oku(yonetici, "konsimento", uzantilar)
    sigorta_metin = dinamik_dosya_bul_ve_oku(yonetici, "sigorta", uzantilar)

    fatura = evrak_sozluge_cevir(fatura_metin)
    konsimento = evrak_sozluge_cevir(konsimento_metin)
    sigorta = evrak_sozluge_cevir(sigorta_metin)
    
    if not kusat_upper:
        print("Hata: Analiz edilecek gelen_kusat dokümanı bulunamadı!")
        sys.exit(1)
        
    t_not, f_not, e_sonuc, z_sonuc, k_tespit, k_eksik, i_not, ucp_sonuc = [], [], [], [], [], [], [], []
    h_var = False
    motor = UCP600KuralMotoru({}, {}, kurallar)

    # 1. TARİH VE VADE SÜRE ANALİZLERİ
    y_match = re.search(r':44C:.*?(\b\d{6}\b)', kusat_upper)
    i_match = re.search(r':48:.*?(\d+)\b', kusat_upper)
    v_match = re.search(r'(\d+)\s*DAYS', kusat_upper)
    bl_tarih_nesnesi = None
    
    if "YUKLEME_TARIHI" in konsimento:
        try:
            bl_tarih_nesnesi = datetime.strptime(konsimento["YUKLEME_TARIHI"].strip(), "%d.%m.%Y")
        except Exception:
            pass

    if y_match:
        y_date = datetime.strptime(y_match.group(1), "%y%m%d")
        t_not.append(f"En Geç Yükleme Tarihi (44C): {y_date.strftime('%d.%m.%Y')}")
        gun = int(i_match.group(1)) if i_match else 21
        son_ibraz = y_date + timedelta(days=gun)
        t_not.append(f"Bankaya Son Evrak İbraz Tarihi: {son_ibraz.strftime('%d.%m.%Y')}")
        
        if bl_tarih_nesnesi and bl_tarih_nesnesi > y_date:
            ucp_sonuc.append(("REZERV RİSKİ", f"Madde 14(c) Geç Yükleme: Konşimento yükleme tarihi ({bl_tarih_nesnesi.strftime('%d.%m.%Y')}), en geç yükleme tarihini ({y_date.strftime('%d.%m.%Y')}) aşmış!"))
            h_var = True

    if v_match and bl_tarih_nesnesi:
        try:
            v_gun = int(v_match.group(1))
            o_t = bl_tarih_nesnesi + timedelta(days=v_gun)
            f_not.append(f"Vade: Konşimentodan {v_gun} gün sonra.")
            f_not.append(f"Tahmini Ödeme Günü: {o_t.strftime('%d.%m.%Y')}")
        except Exception:
            pass

    # 2. INCOTERMS & SİGORTA DENETİMİ
    incoterms = ["EXW", "FCA", "FAS", "FOB", "CFR", "CIF", "CPT", "CIP", "DAP", "DDP"]
    b_in = next((i for i in incoterms if i in kusat_upper), None)
    
    if b_in:
        i_not.append(("BILGI", f"Saptanan teslim sekli: {b_in}"))
        if b_in in ["CIF", "CIP"]:
            if "INSURANCE" in kusat_upper or "110%" in kusat_upper:
                i_not.append(("UYUMLU", "Sigorta sartlari eksiksiz saptandi."))
                if "TUTAR" in fatura and "SIGORTA_TUTARI" in sigorta:
                    try:
                        f_tutar = safe_float(fatura["TUTAR"])
                        s_tutar = safe_float(sigorta["SIGORTA_TUTARI"])
                        if s_tutar < (f_tutar * 1.10):
                            ucp_sonuc.append(("REZERV RİSKİ", f"Madde 28(f)(ii): Sigorta tutarı ({s_tutar:,.2f}), fatura CIF değerinin %110'undan ({(f_tutar*1.10):,.2f}) az!"))
                            h_var = True
                        else:
                            ucp_sonuc.append(("UYUMLU", "Madde 28(f)(ii): Sigorta kapsamı %110 şartını sağlıyor."))
                    except Exception:
                        pass
            else:
                i_not.append(("REZERV RISKI", "Akreditif CIF/CIP olmasına rağmen Sigorta Poliçesi şartı metinde eksik!"))
                h_var = True
    else:
        i_not.append(("REZERV RISKI", "Kritik UYARI: Akreditif metninde resmi bir Incoterms (Teslim Şekli) saptanamadı!"))
        h_var = True

    # 3. UCP 600 MADDE 18 - FATURA VE MAL TANIMI ÇAPRAZ KONTROLLERİ
    m_match = re.search(r':45A:([\s\S]*?)(?=\n:\d{2}[A-Z]?:|$)', kusat_upper)
    if not m_match:
        m_match = re.search(r':46A:.*?COMMERCIAL INVOICE([\s\S]*?)(?=\n:\d{2}[A-Z]?:|$)', kusat_upper)

    if m_match and "MAL_TANIMI" in fatura:
        f_mal_norm = motor.metin_isbp_normalize(fatura["MAL_TANIMI"])
        k_mal_norm = motor.metin_isbp_normalize(m_match.group(1))
        
        if f_mal_norm in k_mal_norm or k_mal_norm in f_mal_norm:
            e_sonuc.append(("UYUMLU", "Mal tanimi fatura ile uyusuyor."))
            ucp_sonuc.append(("UYUMLU", "Madde 18(c): Ticari faturadaki mal tanımı akreditifle uyumlu."))
        else:
            e_sonuc.append(("REZERV RISKI", "Faturadaki mal tanimi kusatla uyusmuyor!"))
            ucp_sonuc.append(("REZERV RİSKİ", "Madde 18(c): Faturadaki mal tanımı akreditifle tam olarak uyuşmuyor!"))
            h_var = True

    if "TARIH" in fatura and bl_tarih_nesnesi:
        try:
            f_d = datetime.strptime(fatura["TARIH"].strip(), "%d.%m.%Y")
            if f_d <= bl_tarih_nesnesi:
                e_sonuc.append(("UYUMLU", "Fatura tarihi gumruk yuklemesinden once."))
            else:
                e_sonuc.append(("REZERV RISKI", "Fatura tarihi yuklemeden sonra olamaz!"))
                h_var = True
        except Exception:
            pass

    if "TARIH" in sigorta and bl_tarih_nesnesi:
        try:
            s_d = datetime.strptime(sigorta["TARIH"].strip(), "%d.%m.%Y")
            if s_d > bl_tarih_nesnesi:
                ucp_sonuc.append(("REZERV RİSKİ", "Madde 28(e): Sigorta poliçesi tarihi yükleme tarihinden sonra olamaz!"))
                h_var = True
            else:
                ucp_sonuc.append(("UYUMLU", "Madde 28(e): Sigorta poliçesi yükleme günü veya öncesinde düzenlenmiş."))
        except Exception:
            pass

    # 4. QUANTITY/AMOUNT TOLERANCE KONTROLÜ
    if "TUTAR" in fatura:
        try:
            f_tutar = safe_float(fatura["TUTAR"])
            akreditif_tutari_match = re.search(r':32B:[A-Z]{3}\s*([\d,.]+)', kusat_upper)
            if akreditif_tutari_match:
                a_tutar = safe_float(akreditif_tutari_match.group(1))
                tolerans_var = "ABOUT" in kusat_upper or "CIRCA" in kusat_upper
                limit = 0.10 if tolerans_var else 0.00
                ust_limit = a_tutar * (1 + limit)
                alt_limit = a_tutar * (1 - limit)
                
                if f_tutar > ust_limit or f_tutar < alt_limit:
                    ucp_sonuc.append(("REZERV RİSKİ", f"Madde 30(a): Fatura tutarı ({f_tutar:,.2f}) tolerans sınırları ({alt_limit:,.2f} - {ust_limit:,.2f}) dışında!"))
                    h_var = True
                else:
                    ucp_sonuc.append(("UYUMLU", "Madde 30(a): Fatura tutarı kabul edilebilir tolerans limitleri içinde."))
        except Exception:
            pass

    # 5 & 6. JSON TABANLI KONTROLLER
    if 'zorunlu_kurallar' in kurallar:
        for k in kurallar['zorunlu_kurallar']:
            k_anahtar_norm = motor.metin_isbp_normalize(k['anahtar'])
            if k_anahtar_norm in motor.metin_isbp_normalize(kusat_upper):
                z_sonuc.append(("UYUMLU", f"'{k['anahtar']}' dogrulandi."))
            else:
                z_sonuc.append(("RISK", f"'{k['anahtar']}' BULUNAMADI!"))
                h_var = True

    if 'kritik_kontroller' in kurallar:
        for k in kurallar['kritik_kontroller']:
            k_anahtar_norm = motor.metin_isbp_normalize(k['anahtar'])
            if k_anahtar_norm in motor.metin_isbp_normalize(kusat_upper):
                k_tespit.append(("TESPIT EDILDI", f"[{k['madde']}] {k['anahtar']} saptandi."))
            else:
                k_eksik.append((k['madde'], k['anahtar']))

    word_raporu_olustur(t_not, z_sonuc, k_tespit, k_eksik, e_sonuc, f_not, i_not, ucp_sonuc, h_var)
    markdown_raporu_olustur(t_not, z_sonuc, k_tespit, k_eksik, e_sonuc, f_not, i_not, ucp_sonuc, h_var)
    print("Yapay Zeka Klasör Destekli Uzman Raporu başarıyla tamamlandı.")

if __name__ == "__main__":
    analiz_yurut()
