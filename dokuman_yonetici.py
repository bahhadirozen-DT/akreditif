import os
import fitz  # PyMuPDF
from docx import Document
import openpyxl
from PIL import Image
import pytesseract

# Eğer Tesseract bilgisayarında özel bir yoldaysa aşağıyı aktifleştirebilirsin:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class DokumanYonetici:
    """Yüklenen tüm dış ticaret evraklarını (Word, Excel, PDF, Resim) 
    otomatik olarak algılayıp temiz metne dönüştüren merkezi modül."""
    
    @staticmethod
    def docx_oku(dosya_yolu):
        """Word dosyasındaki paragrafları ve tabloları okur."""
        doc = Document(dosya_yolu)
        metin_parçalari = []
        
        # 1. Paragrafları oku
        for p in doc.paragraphs:
            if p.text.strip():
                metin_parçalari.append(p.text)
                
        # 2. Tabloların içini oku (Fatura ve çeki listeleri için kritik)
        for tablo in doc.tables:
            for satir in tablo.rows:
                satir_metni = [hucre.text.strip() for hucre in satir.cells if hucre.text.strip()]
                if satir_metni:
                    metin_parçalari.append(" | ".join(satir_metni))
                    
        return "\n".join(metin_parçalari)

    @staticmethod
    def xlsx_oku(dosya_yolu):
        """Excel dosyasındaki tüm satır ve sütunları düz metne çevirir."""
        wb = openpyxl.load_workbook(dosya_yolu, data_only=True)
        metin_parçalari = []
        
        for sheet in wb.worksheets:
            for satir in sheet.iter_rows(values_only=True):
                # Boş olmayan hücreleri stringe çevirip birleştiriyoruz
                satir_filtre = [str(hucre).strip() for hucre in satir if hucre is not None]
                if satir_filtre:
                    metin_parçalari.append(" | ".join(satir_filtre))
                    
        return "\n".join(metin_parçalari)

    @staticmethod
    def pdf_oku(dosya_yolu):
        """PDF dijital ise doğrudan okur, taranmış resim ise OCR tetikler."""
        doc = fitz.open(dosya_yolu)
        tam_metin = ""
        
        for sayfa_no in range(len(doc)):
            sayfa = doc[sayfa_no]
            metin = sayfa.get_text()
            
            # Eğer sayfadan metin geliyorsa dijital PDF'tir
            if metin.strip():
                tam_metin += metin + "\n"
            else:
                # Metin gelmiyorsa bu taranmış bir sayfadır, OCR yapalım
                # Sayfayı resme dönüştür (render)
                pix = sayfa.get_pixmap(dpi=150)
                resim_yolu = f"gecici_sayfa_{sayfa_no}.png"
                pix.save(resim_yolu)
                
                # Resmi OCR ile oku
                ocr_metni = DokumanYonetici.resim_ocr_oku(resim_yolu)
                tam_metin += ocr_metni + "\n"
                
                # Geçici resmi temizle
                if os.path.exists(resim_yolu):
                    os.remove(resim_yolu)
                    
        return tam_metin

    @staticmethod
    def resim_ocr_oku(dosya_yolu):
        """PNG/JPG gibi doğrudan yüklenen tarama resimleri okur."""
        try:
            resim = Image.open(dosya_yolu)
            # Hem Türkçe hem İngilizce karakter desteğiyle oku
            metin = pytesseract.image_to_string(resim, lang='tur+eng')
            return metin
        except Exception as e:
            return f"[OCR HATASI: Tesseract kurulu olmayabilir veya erişilemedi: {str(e)}]"

    def evrak_metne_cevir(self, dosya_yolu):
        """Dosya uzantısına göre otomatik yönlendirme yapar."""
        if not os.path.exists(dosya_yolu):
            raise FileNotFoundError(f"{dosya_yolu} bulunamadı.")
            
        uzanti = os.path.splitext(dosya_yolu)[1].lower()
        
        if uzanti == '.docx':
            return self.docx_oku(dosya_yolu)
        elif uzanti == '.xlsx':
            return self.xlsx_oku(dosya_yolu)
        elif uzanti == '.pdf':
            return self.pdf_oku(dosya_yolu)
        elif uzanti in ['.png', '.jpg', '.jpeg', '.tiff']:
            return self.resim_ocr_oku(dosya_yolu)
        else:
            # Eğer düz metinse (.txt) doğrudan oku
            with open(dosya_yolu, 'r', encoding='utf-8') as f:
                return f.read()

# Test Etmek İçin Küçük Bir Kod Blokçuğu
if __name__ == "__main__":
    yonetici = DokumanYonetici()
    print("Doküman Yönetici Modülü Başarıyla Başlatıldı. Teste Hazır!")
